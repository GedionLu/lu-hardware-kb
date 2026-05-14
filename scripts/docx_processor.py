#!/usr/bin/env python3
"""
DOCX 文档处理工具
python-docx 1.1.2 + Pillow 10.4 + lxml (Python 3.8+)
或 python-docx 0.8.11 (fallback)

功能:
  1. 全文提取（段落、标题、列表、表格）
  2. 图片提取（优先 python-docx >= 1.0 builtin API，降级 zipfile）
  3. 表格结构化提取（转 JSON/CSV）
  4. 文档元数据（标题、作者、页数估算）
  5. 批量处理 + 联合输出

用法:
  python3 docx_processor.py <input.docx> [options]

Options:
  --output-dir DIR      输出目录（默认: ./output）
  --extract-images      提取文档中的图片
  --image-format FORMAT 图片输出格式 (png/jpg) [默认: png]
  --json                输出结构化 JSON（含段落、表格、图片引用）
  --markdown            输出 Markdown 渲染文档
  --batch DIR           批量处理目录下所有 .docx 文件
  --verbose             详细日志
"""

import argparse
import json
import os
import re
import shutil
import sys
import zipfile
from collections import defaultdict
from pathlib import Path
from xml.etree import ElementTree

import docx
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image

# =========== NSMAP for DOCX XML parsing ===========
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
}

XML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def ensure_dir(path):
    Path(path).mkdir(parents=True, exist_ok=True)
    return path


def extract_text(doc: Document) -> list:
    """提取所有段落文本及样式信息"""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()
        if not text:
            continue
        # 检测是否为列表项（通过 numbering）
        is_list = False
        pPr = para._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        if numPr is not None:
            is_list = True

        paragraphs.append(
            {
                "index": i,
                "text": text,
                "style": style_name,
                "is_list": is_list,
                # "runs": [{"text": r.text, "bold": r.bold, "italic": r.italic, "font_size": r.font.size}
                #          for r in para.runs]  # 精确到 run 级别，数据量大时注释掉
            }
        )
    return paragraphs


def extract_tables(doc: Document) -> list:
    """提取所有表格为结构化数据"""
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        # 尝试识别表头（第一行）
        header = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        tables.append(
            {
                "index": t_idx,
                "rows": len(rows),
                "cols": len(header),
                "header": header,
                "data": data_rows,
                "flattened": rows,  # 保留完整数据
                "markdown": table_to_markdown(rows),
            }
        )
    return tables


def table_to_markdown(rows: list) -> str:
    """表格转 Markdown 格式"""
    if not rows:
        return ""
    lines = []
    # 表头
    lines.append("| " + " | ".join(rows[0]) + " |")
    # 分隔线
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # 数据行
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _extract_images_docx_api(doc: Document, output_dir: str) -> list:
    """使用 python-docx >= 1.0 原生 API 提取图片"""
    images = []
    seen = set()
    img_dir = ensure_dir(os.path.join(output_dir, "images"))

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                ext = os.path.splitext(rel.target_ref)[1].lstrip(".") or "png"
                # 生成文件名
                base_name = f"{rel.rId}.{ext}"
                out_path = os.path.join(img_dir, base_name)
                with open(out_path, "wb") as f:
                    f.write(img_data)

                if out_path in seen:
                    continue
                seen.add(out_path)

                # 获取尺寸
                try:
                    with Image.open(out_path) as im:
                        w, h = im.size
                except Exception:
                    w, h = 0, 0

                images.append({
                    "rId": rel.rId,
                    "media_path": rel.target_ref,
                    "saved_path": out_path,
                    "format": ext,
                    "width": w,
                    "height": h,
                    "alt_text": "",
                })
            except Exception:
                continue

    return images


def extract_images_via_zip(docx_path: str, output_dir: str, img_format: str = "png") -> list:
    """
    Fallback：从 DOCX ZIP 结构中提取图片（适用于 python-docx < 1.0）
    通过解析 XML + zipfile 实现。
    """
    images = []
    img_dir = ensure_dir(os.path.join(output_dir, "images"))

    # Step 1: 提取所有 media 文件
    media_map = {}  # {filename_in_zip: saved_path}
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                basename = os.path.basename(name)
                out_path = os.path.join(img_dir, basename)
                # 避免覆盖
                if os.path.exists(out_path):
                    base, ext = os.path.splitext(basename)
                    counter = 1
                    while os.path.exists(out_path):
                        out_path = os.path.join(img_dir, f"{base}_{counter}{ext}")
                        counter += 1
                with open(out_path, "wb") as f:
                    f.write(z.read(name))
                media_map[name] = out_path

    # Step 2: 解析 document.xml 建立图片位置索引
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            xml_content = z.read("word/document.xml")
            doc_xml = etree.fromstring(xml_content)
    except Exception:
        doc_xml = None

    if doc_xml is not None:
        # 获取 rId -> target 映射 (从 word/_rels/document.xml.rels)
        rels_map = {}
        try:
            with zipfile.ZipFile(docx_path, "r") as z:
                rels_content = z.read("word/_rels/document.xml.rels")
                rels_xml = etree.fromstring(rels_content)
                for rel in rels_xml:
                    rid = rel.get("Id")
                    target = rel.get("Target")
                    rel_type = rel.get("Type", "")
                    if rid and target and IMAGE_REL_TYPE in rel_type:
                        # target 可能是相对路径 media/image1.png
                        if not target.startswith("word/"):
                            target = f"word/{target}"
                        rels_map[rid] = target
        except Exception:
            pass

        # 查找所有图片引用并关联位置
        for blip in doc_xml.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
            embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if embed and embed in rels_map:
                media_path = rels_map[embed]
                saved_path = media_map.get(media_path, "")

                if saved_path:
                    # 获取图片尺寸
                    try:
                        with Image.open(saved_path) as img:
                            w, h = img.size
                    except Exception:
                        w, h = 0, 0

                    # 尝试获取可能的位置上下文
                    parent = blip.getparent()
                    while parent is not None:
                        # 往上找段落
                        if parent.tag == f"{{{XML_NS}}}p":
                            break
                        parent = parent.getparent()

                    images.append(
                        {
                            "rId": embed,
                            "media_path": media_path,
                            "saved_path": saved_path,
                            "format": os.path.splitext(saved_path)[1].lstrip("."),
                            "width": w,
                            "height": h,
                            "alt_text": "",  # DOCX 不强制 alt text
                        }
                    )

    # 去重（同一张图片可能在文档中被多次引用）
    seen_paths = set()
    deduped = []
    for img in images:
        if img["saved_path"] not in seen_paths:
            seen_paths.add(img["saved_path"])
            deduped.append(img)
    images = deduped

    # 合并未通过 XML 关联到的 media 文件
    extracted_paths = {img["saved_path"] for img in images}
    for media_path, saved_path in media_map.items():
        if saved_path not in extracted_paths:
            try:
                with Image.open(saved_path) as img:
                    w, h = img.size
            except Exception:
                w, h = 0, 0
            images.append(
                {
                    "rId": "unreferenced",
                    "media_path": media_path,
                    "saved_path": saved_path,
                    "format": os.path.splitext(saved_path)[1].lstrip("."),
                    "width": w,
                    "height": h,
                    "alt_text": "",
                }
            )

    return images


def extract_doc_metadata(doc: Document, docx_path: str) -> dict:
    """提取文档元数据"""
    meta = {
        "file_name": os.path.basename(docx_path),
        "file_size": os.path.getsize(docx_path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
    }

    # Core properties
    props = doc.core_properties
    for attr in ["title", "author", "created", "modified", "subject", "keywords", "category"]:
        val = getattr(props, attr, None)
        if val:
            meta[attr] = str(val)

    return meta


def to_markdown(doc: Document, docx_path: str, images: list = None) -> str:
    """将文档渲染为 Markdown"""
    lines = []

    # 标题
    props = doc.core_properties
    if props.title:
        lines.append(f"# {props.title}\n")

    # 元数据
    meta_lines = []
    if props.author:
        meta_lines.append(f"- **作者**: {props.author}")
    if props.created:
        meta_lines.append(f"- **创建时间**: {props.created}")
    lines.extend(meta_lines)
    if meta_lines:
        lines.append("")

    # 段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else "Normal"

        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("Heading"):
            level = min(6, max(1, int(style.split()[-1])))
            lines.append(f"{'#' * level} {text}")
        elif para._element.find(qn("w:pPr")) is not None:
            pPr = para._element.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
            if numPr is not None:
                lines.append(f"- {text}")
            else:
                lines.append(text)
        else:
            lines.append(text)

        # Runs-level formatting
        for run in para.runs:
            if run.text.strip():
                # Bold / italic markers
                pass  # for now keep simple

    # 表格
    for t_idx, table in enumerate(doc.tables):
        lines.append("")
        lines.append(table_to_markdown([[cell.text.strip() for cell in row.cells] for row in table.rows]))
        lines.append("")

    # 图片引用
    if images:
        lines.append("")
        lines.append("## 文档图片")
        for img in images:
            rel_path = os.path.relpath(img["saved_path"], start=os.path.dirname(docx_path))
            dims = f" ({img['width']}×{img['height']})" if img["width"] else ""
            lines.append(f"![{img['media_path']}]({rel_path}){dims}")
        lines.append("")

    return "\n".join(lines)


def process_single(docx_path: str, output_dir: str, extract_imgs: bool = False,
                   img_format: str = "png", to_json: bool = False, to_md: bool = False,
                   verbose: bool = False) -> dict:
    """处理单个 DOCX 文件"""
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    docx_path = os.path.abspath(docx_path)
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    doc_output = ensure_dir(os.path.join(output_dir, base_name))

    if verbose:
        print(f"处理: {docx_path}")

    doc = Document(docx_path)

    # 1. 元数据
    metadata = extract_doc_metadata(doc, docx_path)
    if verbose:
        print(f"  段落: {metadata['paragraphs']}, 表格: {metadata['tables']}")

    # 2. 提取文本
    paragraphs = extract_text(doc)

    # 3. 提取表格
    tables = extract_tables(doc)

    # 4. 提取图片（自动选择最佳方法）
    images = []
    if extract_imgs:
        # 优先使用 python-docx >= 1.x 原生 API
        docx_version = tuple(int(x) for x in docx.__version__.split("."))
        if docx_version >= (1, 0, 0):
            if verbose:
                print(f"  使用 python-docx {docx.__version__} 原生图片 API")
            images = _extract_images_docx_api(doc, doc_output)
        else:
            if verbose:
                print(f"  使用 ZIP/XML fallback 提取图片")
            images = extract_images_via_zip(docx_path, doc_output, img_format)
        if verbose:
            print(f"  图片: {len(images)} 张")

    # 5. 结构化 JSON 输出
    result = {
        "metadata": metadata,
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
    }

    # 将图片路径改为相对于 JSON 文件的相对路径
    json_dir = doc_output
    for img in result.get("images", []):
        saved = img.get("saved_path", "")
        if saved:
            try:
                rel = os.path.relpath(saved, json_dir)
                img["saved_path"] = rel
            except Exception:
                pass  # 保留绝对路径兜底

    if to_json:
        json_path = os.path.join(doc_output, f"{base_name}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        if verbose:
            print(f"  JSON -> {json_path}")

    if to_md:
        md_content = to_markdown(doc, docx_path, images)
        md_path = os.path.join(doc_output, f"{base_name}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        if verbose:
            print(f"  MD   -> {md_path}")

    return result


def process_batch(batch_dir: str, output_dir: str, extract_imgs: bool = False,
                  img_format: str = "png", to_json: bool = False, to_md: bool = False,
                  verbose: bool = False) -> list:
    """批量处理目录下所有 .docx 文件"""
    docx_files = sorted(Path(batch_dir).glob("*.docx"))
    if not docx_files:
        print(f"未在 {batch_dir} 中找到 .docx 文件")
        return []

    print(f"发现 {len(docx_files)} 个 DOCX 文件")
    results = []
    for f in docx_files:
        try:
            res = process_single(
                str(f), output_dir, extract_imgs, img_format, to_json, to_md, verbose
            )
            results.append(res)
            print(f"  OK: {f.name}")
        except Exception as e:
            print(f"  FAIL: {f.name} -> {e}")

    # 批量汇总
    summary_path = os.path.join(output_dir, "_batch_summary.json")
    summary = {
        "total": len(docx_files),
        "success": len(results),
        "failed": len(docx_files) - len(results),
        "files": [
            {
                "file": r["metadata"]["file_name"],
                "paragraphs": r["metadata"]["paragraphs"],
                "tables": r["metadata"]["tables"],
                "images": len(r["images"]),
            }
            for r in results
        ],
    }
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"\n批量汇总 -> {summary_path}")
    return results


# =========== CLI ===========
def main():
    parser = argparse.ArgumentParser(description="DOCX 文档处理工具")
    parser.add_argument("input", help="输入 DOCX 文件或目录（配合 --batch）")
    parser.add_argument("--output-dir", default="./output", help="输出目录")
    parser.add_argument("--extract-images", action="store_true", help="提取图片")
    parser.add_argument("--image-format", default="png", choices=["png", "jpg"], help="图片输出格式")
    parser.add_argument("--json", action="store_true", help="输出 JSON")
    parser.add_argument("--markdown", action="store_true", help="输出 Markdown")
    parser.add_argument("--batch", action="store_true", help="批量模式")
    parser.add_argument("--verbose", "-v", action="store_true", help="详细输出")
    args = parser.parse_args()

    if args.batch:
        process_batch(args.input, args.output_dir, args.extract_images,
                      args.image_format, args.json, args.markdown, args.verbose)
    else:
        process_single(args.input, args.output_dir, args.extract_images,
                       args.image_format, args.json, args.markdown, args.verbose)
        print(f"\n输出目录: {os.path.abspath(args.output_dir)}")


if __name__ == "__main__":
    main()
